#!/usr/bin/env python3
"""회의록 마크다운을 한국식 회의록 .docx(단일 표 양식)로 변환한다.

ddokham-meeting-minute 스킬의 양식을 python-docx 로 재현:
- 전체를 하나의 표로 구성
- 제목행/내용 머리행은 연청색 배경·굵게·가운데
- 일시/장소, 주관/작성자, 참석자/보고대상 행
- 내용 본문은 ㅇ 주 항목, - 하위 항목

사용: python scripts/md_to_docx.py <회의록.md> [출력.docx]
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

BLUE = "D6EAF8"
FONT = "맑은 고딕"


def parse_md(md_text):
    """회의록 md 에서 구조를 추출."""
    data = {"title": "", "일시": "", "장소": "", "참석자": "", "작성일": "",
            "sections": []}
    lines = md_text.splitlines()
    # 제목
    for ln in lines:
        m = re.match(r"#\s*회의록\s*[—-]\s*(.+)", ln)
        if m:
            data["title"] = m.group(1).strip()
            break
    # 헤더 표 (| 항목 | 내용 |)
    for ln in lines:
        m = re.match(r"\|\s*(일시|장소|참석자|작성일)\s*\|\s*(.+?)\s*\|", ln)
        if m:
            data[m.group(1)] = m.group(2).strip()
    # 본문 섹션 (## 로 시작)
    cur = None
    for ln in lines:
        h = re.match(r"##\s+(.*)", ln)
        if h:
            cur = {"head": h.group(1).strip(), "body": []}
            data["sections"].append(cur)
        elif cur is not None:
            cur["body"].append(ln)
    return data


def shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), color)
    tcPr.append(sh)


def set_cell(cell, text, bold=False, center=False, size=11,
             valign=WD_ALIGN_VERTICAL.CENTER):
    cell.vertical_alignment = valign
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold = bold
    run.font.name = FONT
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    return cell


def add_body_paragraph(cell, text, first=False):
    para = cell.paragraphs[0] if first else cell.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(11)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def build_content_lines(data):
    """본문 섹션을 ㅇ/- 글머리 텍스트 라인으로 변환."""
    out = []
    for sec in data["sections"]:
        out.append(("ㅇ", sec["head"]))
        for raw in sec["body"]:
            s = raw.strip()
            if not s:
                continue
            if s.startswith("###"):
                out.append(("-", s.lstrip("#").strip()))
            elif s.startswith(("- ", "* ")):
                out.append(("  -", s[2:].strip()))
            elif re.match(r"^\d+\.\s", s):
                out.append(("  -", re.sub(r"^\d+\.\s", "", s)))
            elif s.startswith("|"):
                # 표 행 → 파이프 정리해서 한 줄로
                cells = [c.strip() for c in s.strip("|").split("|")]
                if set("".join(cells)) <= set("-: "):
                    continue  # 구분선
                out.append(("  -", " / ".join(c for c in cells if c)))
            else:
                out.append(("  ", s))
    return out


def convert(md_path, out_path):
    data = parse_md(Path(md_path).read_text(encoding="utf-8"))
    doc = Document()
    doc.styles["Normal"].font.name = FONT
    doc.styles["Normal"].element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    table = doc.add_table(rows=0, cols=3)
    table.style = "Table Grid"

    # Row 1: 제목
    r = table.add_row().cells
    a = r[0].merge(r[1]).merge(r[2])
    set_cell(a, "회 의 록", bold=True, center=True, size=15)
    shade(a, BLUE)

    # Row 2: 일시 / 장소
    r = table.add_row().cells
    left = r[0].merge(r[1])
    set_cell(left, f"일 시 : {data.get('일시','')}")
    set_cell(r[2], f"회의장소 : {data.get('장소','') or '(미정)'}")

    # Row 3: 주관 / 작성자
    r = table.add_row().cells
    left = r[0].merge(r[1])
    set_cell(left, "주 관 : (미정)")
    set_cell(r[2], f"작 성 자 : 회의록 자동화")

    # Row 4: 참석자 / 보고대상
    r = table.add_row().cells
    left = r[0].merge(r[1])
    set_cell(left, f"참 석 자 : {data.get('참석자','')}")
    set_cell(r[2], "보고대상 : 팀원 전체")

    # Row 5: 내용 머리행
    r = table.add_row().cells
    a = r[0].merge(r[1]).merge(r[2])
    set_cell(a, "내 용", bold=True, center=True, size=12)
    shade(a, BLUE)

    # Row 6: 내용 본문
    r = table.add_row().cells
    body = r[0].merge(r[1]).merge(r[2])
    body.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    lines = build_content_lines(data)
    first = True
    for marker, text in lines:
        add_body_paragraph(body, f"{marker} {text}".rstrip(), first=first)
        first = False

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"[성공] docx 생성: {out_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        sys.exit("사용: python scripts/md_to_docx.py <회의록.md> [출력.docx]")
    md = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else str(Path(md).with_suffix(".docx"))
    convert(md, out)
